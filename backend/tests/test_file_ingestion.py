import asyncio
import io
import tempfile
import unittest
import zipfile
from pathlib import Path

from fastapi import UploadFile
from openpyxl import Workbook

from config import settings
from services.file_ingestion import (
    FileIngestionError,
    analyze_uploaded_document,
    build_file_agent_prompt,
    save_and_analyze_uploads,
    sanitize_upload_filename,
)
from services.project_files import scan_task_workspace
from tools.tool_executor import _requested_document_artifact_error


def upload_file(name: str, data: bytes) -> UploadFile:
    return UploadFile(filename=name, file=io.BytesIO(data))


def xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Dados"
    sheet.append(["Item", "Valor"])
    sheet.append(["Receita", 1200])
    sheet.append(["Custo", 800])
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def zip_bytes(entries: dict[str, bytes]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in entries.items():
            archive.writestr(name, data)
    return buffer.getvalue()


class FileIngestionTests(unittest.TestCase):
    def test_sanitizes_names_and_rejects_unsupported_extensions(self) -> None:
        self.assertEqual(sanitize_upload_filename("../Meu Arquivo.xlsx"), "Meu_Arquivo.xlsx")
        self.assertEqual(sanitize_upload_filename("dados.zip"), "dados.zip")
        with self.assertRaises(FileIngestionError):
            sanitize_upload_filename("script.exe")

    def test_extracts_xlsx_csv_docx_txt_md_json_and_pdf_metadata(self) -> None:
        from docx import Document

        previous_workspace = settings.WORKSPACE_PATH
        with tempfile.TemporaryDirectory() as tmp:
            try:
                settings.WORKSPACE_PATH = Path(tmp)
                root = settings.WORKSPACE_PATH / "task"
                root.mkdir()
                xlsx_path = root / "dados.xlsx"
                xlsx_path.write_bytes(xlsx_bytes())
                csv_path = root / "dados.csv"
                csv_path.write_text("item,valor\nreceita,1200\ncusto,800\n", encoding="utf-8")
                docx_path = root / "relatorio.docx"
                document = Document()
                document.add_heading("Relatorio", level=1)
                document.add_paragraph("Conteudo do arquivo para analise.")
                document.save(docx_path)
                txt_path = root / "notas.txt"
                txt_path.write_text("Linha importante\nOutra linha\n", encoding="utf-8")
                md_path = root / "guia.md"
                md_path.write_text("# Guia\n\nConteudo em markdown.\n", encoding="utf-8")
                json_path = root / "dados.json"
                json_path.write_text('{"cliente": "ACME", "total": 1200}', encoding="utf-8")

                self.assertEqual(analyze_uploaded_document("task", xlsx_path)["kind"], "spreadsheet")
                self.assertEqual(analyze_uploaded_document("task", csv_path)["kind"], "csv")
                self.assertEqual(analyze_uploaded_document("task", docx_path)["kind"], "word")
                self.assertEqual(analyze_uploaded_document("task", txt_path)["kind"], "text")
                self.assertEqual(analyze_uploaded_document("task", md_path)["kind"], "markdown")
                self.assertTrue(analyze_uploaded_document("task", json_path)["details"]["valid"])

                try:
                    from pypdf import PdfWriter
                except ImportError:
                    self.skipTest("pypdf nao instalado")
                pdf_path = root / "vazio.pdf"
                writer = PdfWriter()
                writer.add_blank_page(width=72, height=72)
                with pdf_path.open("wb") as handle:
                    writer.write(handle)
                self.assertEqual(analyze_uploaded_document("task", pdf_path)["kind"], "pdf")
            finally:
                settings.WORKSPACE_PATH = previous_workspace

    def test_uploads_are_preserved_and_do_not_satisfy_edited_artifact_gate(self) -> None:
        previous_workspace = settings.WORKSPACE_PATH
        with tempfile.TemporaryDirectory() as tmp:
            try:
                settings.WORKSPACE_PATH = Path(tmp)
                task_id = "task-files"
                analyses = asyncio.run(
                    save_and_analyze_uploads(
                        task_id,
                        [upload_file("entrada.xlsx", xlsx_bytes())],
                    )
                )
                self.assertEqual(analyses[0]["path"], "uploads/entrada.xlsx")
                self.assertEqual(scan_task_workspace(settings.WORKSPACE_PATH / task_id), [])

                missing = _requested_document_artifact_error(task_id, "vertex 'edite essa planilha'")
                self.assertIn(".xlsx real", missing)

                output_dir = settings.WORKSPACE_PATH / task_id / "outputs"
                output_dir.mkdir()
                (output_dir / "entrada_editada.xlsx").write_bytes(xlsx_bytes())
                complete = _requested_document_artifact_error(task_id, "vertex 'edite essa planilha'")
            finally:
                settings.WORKSPACE_PATH = previous_workspace

        self.assertIsNone(complete)

    def test_zip_is_extracted_and_inner_supported_files_are_analyzed(self) -> None:
        previous_workspace = settings.WORKSPACE_PATH
        with tempfile.TemporaryDirectory() as tmp:
            try:
                settings.WORKSPACE_PATH = Path(tmp)
                analyses = asyncio.run(
                    save_and_analyze_uploads(
                        "task-zip",
                        [
                            upload_file(
                                "pacote.zip",
                                zip_bytes(
                                    {
                                        "docs/notas.txt": b"Resumo do pacote",
                                        "dados/planilha.xlsx": xlsx_bytes(),
                                    }
                                ),
                            )
                        ],
                    )
                )
            finally:
                settings.WORKSPACE_PATH = previous_workspace

        archive = next(item for item in analyses if item["kind"] == "archive")
        self.assertEqual(archive["path"], "uploads/pacote.zip")
        self.assertEqual(archive["details"]["file_count"], 2)
        self.assertTrue(any(item["path"].endswith("docs/notas.txt") for item in analyses))
        self.assertTrue(any(item["kind"] == "spreadsheet" for item in analyses))

    def test_zip_rejects_path_traversal(self) -> None:
        previous_workspace = settings.WORKSPACE_PATH
        with tempfile.TemporaryDirectory() as tmp:
            try:
                settings.WORKSPACE_PATH = Path(tmp)
                with self.assertRaises(FileIngestionError):
                    asyncio.run(
                        save_and_analyze_uploads(
                            "task-zip-bad",
                            [upload_file("bad.zip", zip_bytes({"../escape.txt": b"no"}))],
                        )
                    )
            finally:
                settings.WORKSPACE_PATH = previous_workspace

    def test_zip_with_shapefile_is_grouped_and_analyzed(self) -> None:
        try:
            import geopandas as gpd
            from shapely.geometry import Point
        except ImportError:
            self.skipTest("dependencias geoespaciais nao instaladas")

        previous_workspace = settings.WORKSPACE_PATH
        with tempfile.TemporaryDirectory() as tmp:
            try:
                settings.WORKSPACE_PATH = Path(tmp)
                source_dir = Path(tmp) / "source"
                source_dir.mkdir()
                gdf = gpd.GeoDataFrame(
                    {"nome": ["A", "B"], "valor": [10, 20]},
                    geometry=[Point(-46.63, -23.55), Point(-46.64, -23.56)],
                    crs="EPSG:4326",
                )
                shp_path = source_dir / "pontos.shp"
                gdf.to_file(shp_path, driver="ESRI Shapefile")
                entries = {
                    f"geo/{path.name}": path.read_bytes()
                    for path in source_dir.iterdir()
                    if path.is_file()
                }

                analyses = asyncio.run(
                    save_and_analyze_uploads(
                        "task-shp",
                        [upload_file("geo.zip", zip_bytes(entries))],
                    )
                )
            finally:
                settings.WORKSPACE_PATH = previous_workspace

        layer = next(item for item in analyses if item["kind"] == "shapefile")
        self.assertTrue(layer["details"]["complete"])
        self.assertEqual(layer["details"]["feature_count"], 2)
        self.assertIn("Point", layer["details"]["geometry_types"])
        self.assertIn("EPSG", layer["details"]["crs"])

    def test_incomplete_shapefile_reports_clear_warning(self) -> None:
        previous_workspace = settings.WORKSPACE_PATH
        with tempfile.TemporaryDirectory() as tmp:
            try:
                settings.WORKSPACE_PATH = Path(tmp)
                analyses = asyncio.run(
                    save_and_analyze_uploads(
                        "task-shp-missing",
                        [upload_file("geo.zip", zip_bytes({"geo/layer.shp": b"fake"}))],
                    )
                )
            finally:
                settings.WORKSPACE_PATH = previous_workspace

        layer = next(item for item in analyses if item["kind"] == "shapefile")
        self.assertFalse(layer["details"]["complete"])
        self.assertIn(".dbf", layer["summary"])

    def test_agent_prompt_keeps_original_uploads_read_only(self) -> None:
        prompt = build_file_agent_prompt(
            "edite essa planilha",
            [
                {
                    "path": "uploads/entrada.xlsx",
                    "extension": ".xlsx",
                    "size_bytes": 123,
                    "summary": "Excel com dados.",
                    "extracted_text": "Item | Valor",
                }
            ],
        )

        self.assertIn("Nao sobrescreva arquivos em uploads/", prompt)
        self.assertIn("gere uma nova versao real em outputs/", prompt)

    def test_agent_prompt_instructs_geospatial_workflow(self) -> None:
        prompt = build_file_agent_prompt(
            "edite a tabela de atributos desse shapefile",
            [
                {
                    "path": "uploads/geo.zip",
                    "extension": ".zip",
                    "size_bytes": 123,
                    "summary": "ZIP com shapefile.",
                    "extracted_text": "",
                    "archive_contents": [{"path": "uploads/archives/geo/layer.shp", "extension": ".shp", "size_bytes": 10}],
                    "geospatial_layers": [
                        {
                            "path": "uploads/archives/geo/layer.shp",
                            "summary": "Shapefile com 2 feicoes.",
                            "details": {"crs": "EPSG:4326", "bounds": [-1, -1, 1, 1]},
                        }
                    ],
                }
            ],
        )

        self.assertIn("CAMADAS_GEOESPACIAIS_DETECTADAS", prompt)
        self.assertIn("ENTREGA_OBRIGATORIA", prompt)
        self.assertIn("geopandas", prompt)
        self.assertIn("nao entregue apenas .shp isolado", prompt)


if __name__ == "__main__":
    unittest.main()
