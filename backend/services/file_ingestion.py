from __future__ import annotations

import csv
import math
import json
import re
import zipfile
from pathlib import Path
from typing import Any

from fastapi import UploadFile
from openpyxl import load_workbook

from config import settings
from services.document_intent import geospatial_delivery_from_text


ALLOWED_DOCUMENT_EXTENSIONS = {".xlsx", ".csv", ".docx", ".pdf", ".txt", ".md", ".json", ".zip"}
ANALYZABLE_DOCUMENT_EXTENSIONS = {".xlsx", ".csv", ".docx", ".pdf", ".txt", ".md", ".json"}
SHAPEFILE_EXTENSIONS = {".shp", ".shx", ".dbf", ".prj", ".cpg"}
SHAPEFILE_REQUIRED_EXTENSIONS = {".shp", ".shx", ".dbf"}
MAX_DOCUMENT_UPLOAD_BYTES = 25 * 1024 * 1024
MAX_ZIP_UPLOAD_BYTES = 100 * 1024 * 1024
MAX_ZIP_EXTRACTED_BYTES = 250 * 1024 * 1024
MAX_ZIP_FILES = 500
TEXT_PREVIEW_CHARS = 8000
PROMPT_PREVIEW_CHARS = 2200
GEOSPATIAL_EDIT_HINT_RE = re.compile(
    r"\b(edite|editar|altere|alterar|atualize|atualizar|corrija|corrigir|"
    r"adicione|adicionar|remova|remover|mude|mudar|campo|campos|coluna|colunas|"
    r"atributo|atributos|tabela\s+de\s+atributos|converter|converta|exportar|exporte|"
    r"zip|compacte)\b",
    re.IGNORECASE,
)


class FileIngestionError(ValueError):
    pass


def _workspace(task_id: str) -> Path:
    return settings.WORKSPACE_PATH / task_id


def uploads_dir(task_id: str) -> Path:
    return _workspace(task_id) / "uploads"


def _relative(task_id: str, path: Path) -> str:
    return str(path.relative_to(_workspace(task_id))).replace("\\", "/")


def sanitize_upload_filename(filename: str | None) -> str:
    raw = Path(str(filename or "arquivo")).name.strip()
    if not raw or raw in {".", ".."}:
        raw = "arquivo"
    raw = raw.replace("\x00", "")
    stem = Path(raw).stem.strip() or "arquivo"
    suffix = Path(raw).suffix.lower()
    if suffix not in ALLOWED_DOCUMENT_EXTENSIONS:
        raise FileIngestionError(f"Extensao nao suportada: {suffix or 'sem extensao'}")
    safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._-") or "arquivo"
    return f"{safe_stem[:90]}{suffix}"


def _unique_path(directory: Path, filename: str) -> Path:
    candidate = directory / filename
    if not candidate.exists():
        return candidate
    stem = candidate.stem
    suffix = candidate.suffix
    for index in range(2, 1000):
        next_candidate = directory / f"{stem}_{index}{suffix}"
        if not next_candidate.exists():
            return next_candidate
    raise FileIngestionError("Nao foi possivel gerar um nome unico para o arquivo.")


def _safe_extract_member_path(base_dir: Path, member_name: str) -> Path:
    normalized = str(member_name or "").replace("\\", "/").strip()
    if not normalized or normalized.endswith("/"):
        raise FileIngestionError("Entrada vazia no ZIP.")
    member_path = Path(normalized)
    if member_path.is_absolute() or any(part in {"", ".", ".."} for part in member_path.parts):
        raise FileIngestionError(f"Caminho inseguro dentro do ZIP: {member_name}")
    target = (base_dir / member_path).resolve()
    base = base_dir.resolve()
    if target != base and base not in target.parents:
        raise FileIngestionError(f"Caminho fora da pasta de extracao: {member_name}")
    return target


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8"):
        try:
            return path.read_text(encoding=encoding, errors="strict")
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _clip(text: str, limit: int = TEXT_PREVIEW_CHARS) -> str:
    value = str(text or "").strip()
    if len(value) <= limit:
        return value
    return value[:limit].rstrip() + "\n...[truncado]"


def _stringify_cell(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _analyze_xlsx(path: Path) -> dict[str, Any]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    formula_workbook = load_workbook(path, read_only=True, data_only=False)
    try:
        sheets: list[dict[str, Any]] = []
        text_lines: list[str] = []
        for worksheet in workbook.worksheets[:8]:
            rows = []
            populated_cells = 0
            header: list[str] = []
            for row in worksheet.iter_rows(max_row=30, values_only=True):
                values = [_stringify_cell(value) for value in row]
                if any(values):
                    populated_cells += sum(1 for value in values if value)
                    if not header:
                        header = values[:20]
                    if len(rows) < 8:
                        rows.append(values[:20])
            formula_count = 0
            formula_sheet = formula_workbook[worksheet.title]
            for row in formula_sheet.iter_rows(max_row=200):
                for cell in row:
                    value = cell.value
                    if isinstance(value, str) and value.startswith("="):
                        formula_count += 1
            sheets.append(
                {
                    "name": worksheet.title,
                    "max_row": worksheet.max_row,
                    "max_column": worksheet.max_column,
                    "header": header,
                    "sample_rows": rows,
                    "populated_cells_sample": populated_cells,
                    "formula_count_sample": formula_count,
                }
            )
            text_lines.append(
                f"Aba {worksheet.title}: {worksheet.max_row} linhas x {worksheet.max_column} colunas; "
                f"cabecalho: {', '.join(item for item in header if item) or 'nao identificado'}"
            )
            for sample in rows[:5]:
                text_lines.append(" | ".join(sample))
        return {
            "kind": "spreadsheet",
            "summary": f"Excel com {len(workbook.worksheets)} aba(s). " + " ".join(text_lines[:8]),
            "extracted_text": _clip("\n".join(text_lines)),
            "details": {"sheets": sheets},
        }
    finally:
        workbook.close()
        formula_workbook.close()


def _analyze_csv(path: Path) -> dict[str, Any]:
    text = _read_text(path)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except csv.Error:
        dialect = csv.excel
    rows: list[list[str]] = []
    for row in csv.reader(text.splitlines(), dialect):
        if any(str(cell).strip() for cell in row):
            rows.append([str(cell).strip() for cell in row])
        if len(rows) >= 30:
            break
    header = rows[0] if rows else []
    max_columns = max((len(row) for row in rows), default=0)
    return {
        "kind": "csv",
        "summary": f"CSV com amostra de {len(rows)} linha(s), {max_columns} coluna(s). Cabecalho: {', '.join(header[:12]) or 'nao identificado'}.",
        "extracted_text": _clip("\n".join(",".join(row) for row in rows)),
        "details": {"header": header, "sample_rows": rows[:10], "columns": max_columns},
    }


def _analyze_docx(path: Path) -> dict[str, Any]:
    from docx import Document

    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines: list[str] = []
    for table in document.tables[:6]:
        for row in table.rows[:8]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                table_lines.append(" | ".join(cells))
    text = "\n".join([*paragraphs[:80], *table_lines[:40]])
    return {
        "kind": "word",
        "summary": f"DOCX com {len(paragraphs)} paragrafo(s) e {len(document.tables)} tabela(s).",
        "extracted_text": _clip(text),
        "details": {"paragraph_count": len(paragraphs), "table_count": len(document.tables)},
    }


def _analyze_pdf(path: Path) -> dict[str, Any]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise FileIngestionError("Leitura de PDF requer a dependencia pypdf instalada.") from exc

    reader = PdfReader(str(path))
    page_texts = []
    for index, page in enumerate(reader.pages[:8], start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            page_texts.append(f"Pagina {index}:\n{text.strip()}")
    extracted = "\n\n".join(page_texts)
    return {
        "kind": "pdf",
        "summary": f"PDF com {len(reader.pages)} pagina(s). Texto extraido em {len(page_texts)} pagina(s) da amostra.",
        "extracted_text": _clip(extracted),
        "details": {"page_count": len(reader.pages), "sample_text_pages": len(page_texts)},
    }


def _analyze_json(path: Path) -> dict[str, Any]:
    text = _read_text(path)
    try:
        data = json.loads(text)
    except json.JSONDecodeError as exc:
        return {
            "kind": "json",
            "summary": f"JSON invalido: {exc.msg} na linha {exc.lineno}, coluna {exc.colno}.",
            "extracted_text": _clip(text),
            "details": {"valid": False, "error": str(exc)},
        }
    if isinstance(data, dict):
        shape = f"objeto com chaves: {', '.join(list(data.keys())[:20])}"
    elif isinstance(data, list):
        shape = f"lista com {len(data)} item(ns)"
    else:
        shape = type(data).__name__
    return {
        "kind": "json",
        "summary": f"JSON valido: {shape}.",
        "extracted_text": _clip(json.dumps(data, ensure_ascii=False, indent=2)),
        "details": {"valid": True, "shape": shape},
    }


def _analyze_text(path: Path, kind: str) -> dict[str, Any]:
    text = _read_text(path)
    lines = [line for line in text.splitlines() if line.strip()]
    return {
        "kind": kind,
        "summary": f"Arquivo {path.suffix.lower()} com {len(lines)} linha(s) nao vazia(s) e {len(text)} caractere(s).",
        "extracted_text": _clip(text),
        "details": {"line_count": len(lines), "char_count": len(text)},
    }


def _json_safe(value: Any) -> Any:
    if isinstance(value, float):
        if math.isfinite(value):
            return value
        return None
    if isinstance(value, dict):
        return {str(key): _json_safe(item) for key, item in value.items()}
    if isinstance(value, list):
        return [_json_safe(item) for item in value]
    return value


def _shape_group_key(path: Path) -> str:
    return str(path.with_suffix("")).replace("\\", "/")


def _collect_shapefile_groups(root: Path) -> dict[str, dict[str, Path]]:
    groups: dict[str, dict[str, Path]] = {}
    for path in sorted(root.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SHAPEFILE_EXTENSIONS:
            continue
        key = _shape_group_key(path)
        groups.setdefault(key, {})[path.suffix.lower()] = path
    return groups


def _analyze_shapefile(task_id: str, components: dict[str, Path]) -> dict[str, Any]:
    missing = sorted(SHAPEFILE_REQUIRED_EXTENSIONS - set(components))
    shp_path = components.get(".shp")
    representative = shp_path or next(iter(components.values()))
    base = {
        "path": _relative(task_id, representative),
        "name": representative.with_suffix("").name,
        "extension": ".shp",
        "size_bytes": sum(path.stat().st_size for path in components.values() if path.exists()),
        "kind": "shapefile",
        "component_paths": [_relative(task_id, path) for path in sorted(components.values())],
    }
    if missing:
        return {
            **base,
            "summary": "Shapefile incompleto; faltam componentes obrigatorios: " + ", ".join(missing),
            "extracted_text": "",
            "details": {"complete": False, "missing_components": missing},
        }

    try:
        import geopandas as gpd
    except ImportError as exc:
        raise FileIngestionError("Analise de shapefile requer geopandas/pyogrio/shapely/pyproj instalados.") from exc

    try:
        gdf = gpd.read_file(shp_path)
    except Exception as exc:
        raise FileIngestionError(f"Nao foi possivel ler shapefile {representative.name}: {type(exc).__name__}") from exc

    attribute_columns = [column for column in gdf.columns if column != gdf.geometry.name]
    geometry_types = sorted(str(item) for item in gdf.geom_type.dropna().unique().tolist())
    crs = str(gdf.crs) if gdf.crs is not None else ""
    bounds = [float(item) for item in gdf.total_bounds.tolist()] if len(gdf) else []
    invalid_count = int((~gdf.geometry.is_valid).sum()) if len(gdf) and gdf.geometry.name in gdf else 0
    sample_rows = [
        _json_safe(row)
        for row in gdf[attribute_columns].head(8).fillna("").astype(str).to_dict(orient="records")
    ]
    column_types = {column: str(dtype) for column, dtype in gdf[attribute_columns].dtypes.items()}
    projected = bool(gdf.crs and getattr(gdf.crs, "is_projected", False))
    area_sum = float(gdf.geometry.area.sum()) if projected and len(gdf) else None
    length_sum = float(gdf.geometry.length.sum()) if projected and len(gdf) else None
    warnings: list[str] = []
    if not crs:
        warnings.append("CRS ausente; nao assuma projecao, area ou distancia sem confirmar.")
    if invalid_count:
        warnings.append(f"{invalid_count} geometria(s) invalida(s) detectada(s).")

    text_lines = [
        f"Layer {representative.with_suffix('').name}",
        f"Feicoes: {len(gdf)}",
        f"CRS: {crs or 'ausente'}",
        f"Geometrias: {', '.join(geometry_types) or 'nao identificado'}",
        f"Bounds: {bounds or 'sem bounds'}",
        f"Colunas: {', '.join(attribute_columns[:30]) or 'sem atributos'}",
    ]
    if warnings:
        text_lines.append("Avisos: " + " ".join(warnings))
    for row in sample_rows:
        text_lines.append(json.dumps(row, ensure_ascii=False))

    summary = (
        f"Shapefile com {len(gdf)} feicao(oes), geometria(s) {', '.join(geometry_types) or 'nao identificada'}, "
        f"CRS {crs or 'ausente'}, {len(attribute_columns)} coluna(s) de atributos."
    )
    return {
        **base,
        "summary": summary,
        "extracted_text": _clip("\n".join(text_lines)),
        "details": {
            "complete": True,
            "feature_count": len(gdf),
            "crs": crs,
            "geometry_types": geometry_types,
            "bounds": bounds,
            "columns": attribute_columns,
            "column_types": column_types,
            "sample_rows": sample_rows,
            "invalid_geometry_count": invalid_count,
            "area_sum": area_sum,
            "length_sum": length_sum,
            "warnings": warnings,
        },
    }


def _analyze_supported_tree(task_id: str, root: Path, *, source_archive: str = "") -> list[dict[str, Any]]:
    analyses: list[dict[str, Any]] = []
    shapefile_component_paths: set[Path] = set()
    for components in _collect_shapefile_groups(root).values():
        shapefile_component_paths.update(components.values())
        analysis = _analyze_shapefile(task_id, components)
        if source_archive:
            analysis["source_archive"] = source_archive
        analyses.append(analysis)

    for path in sorted(root.rglob("*")):
        if not path.is_file() or path in shapefile_component_paths:
            continue
        suffix = path.suffix.lower()
        if suffix not in ANALYZABLE_DOCUMENT_EXTENSIONS:
            continue
        try:
            analysis = analyze_uploaded_document(task_id, path)
        except FileIngestionError as exc:
            analysis = {
                "path": _relative(task_id, path),
                "name": path.name,
                "extension": suffix,
                "size_bytes": path.stat().st_size,
                "kind": "document_error",
                "summary": str(exc),
                "extracted_text": "",
                "details": {"error": str(exc)},
            }
        if source_archive:
            analysis["source_archive"] = source_archive
        analyses.append(analysis)
    return analyses


def _analyze_zip(task_id: str, zip_path: Path) -> list[dict[str, Any]]:
    archive_root = uploads_dir(task_id) / "archives"
    archive_root.mkdir(parents=True, exist_ok=True)
    extract_dir = _unique_path(archive_root, zip_path.stem)
    extract_dir.mkdir(parents=True, exist_ok=False)

    try:
        with zipfile.ZipFile(zip_path) as archive:
            members = [item for item in archive.infolist() if not item.is_dir()]
            if len(members) > MAX_ZIP_FILES:
                raise FileIngestionError(f"ZIP contem {len(members)} arquivos; limite e {MAX_ZIP_FILES}.")
            extracted_total = 0
            extracted_files: list[dict[str, Any]] = []
            for member in members:
                target = _safe_extract_member_path(extract_dir, member.filename)
                size = int(member.file_size or 0)
                if size <= 0:
                    raise FileIngestionError(f"Arquivo vazio dentro do ZIP: {member.filename}")
                extracted_total += size
                if extracted_total > MAX_ZIP_EXTRACTED_BYTES:
                    raise FileIngestionError("Conteudo extraido do ZIP excede 250 MB.")
                target.parent.mkdir(parents=True, exist_ok=True)
                with archive.open(member) as source, target.open("wb") as destination:
                    while True:
                        chunk = source.read(1024 * 1024)
                        if not chunk:
                            break
                        destination.write(chunk)
                extracted_files.append(
                    {
                        "path": _relative(task_id, target),
                        "name": target.name,
                        "extension": target.suffix.lower(),
                        "size_bytes": size,
                    }
                )
    except zipfile.BadZipFile as exc:
        raise FileIngestionError("ZIP invalido ou corrompido.") from exc

    inner_analyses = _analyze_supported_tree(task_id, extract_dir, source_archive=_relative(task_id, zip_path))
    layers = [item for item in inner_analyses if item.get("kind") == "shapefile"]
    archive_analysis = {
        "path": _relative(task_id, zip_path),
        "name": zip_path.name,
        "extension": ".zip",
        "size_bytes": zip_path.stat().st_size,
        "kind": "archive",
        "summary": (
            f"ZIP extraido com {len(extracted_files)} arquivo(s), "
            f"{len(inner_analyses)} item(ns) analisavel(is) e {len(layers)} layer(s) geoespacial(is)."
        ),
        "extracted_text": "\n".join(
            f"- {item['path']} ({item.get('size_bytes', 0)} bytes)" for item in extracted_files[:120]
        ),
        "details": {
            "extract_dir": _relative(task_id, extract_dir),
            "file_count": len(extracted_files),
            "total_uncompressed_size": sum(int(item.get("size_bytes") or 0) for item in extracted_files),
        },
        "archive_contents": extracted_files[:500],
        "geospatial_layers": [
            {
                "path": item.get("path"),
                "name": item.get("name"),
                "summary": item.get("summary"),
                "details": item.get("details", {}),
            }
            for item in layers
        ],
    }
    return [archive_analysis, *inner_analyses]


def analyze_uploaded_document(task_id: str, path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    base: dict[str, Any] = {
        "path": _relative(task_id, path),
        "name": path.name,
        "extension": suffix,
        "size_bytes": path.stat().st_size,
    }
    if suffix == ".xlsx":
        analysis = _analyze_xlsx(path)
    elif suffix == ".csv":
        analysis = _analyze_csv(path)
    elif suffix == ".docx":
        analysis = _analyze_docx(path)
    elif suffix == ".pdf":
        analysis = _analyze_pdf(path)
    elif suffix == ".json":
        analysis = _analyze_json(path)
    elif suffix == ".md":
        analysis = _analyze_text(path, "markdown")
    elif suffix == ".txt":
        analysis = _analyze_text(path, "text")
    else:
        raise FileIngestionError(f"Extensao nao suportada: {suffix}")
    return {**base, **analysis}


async def save_and_analyze_uploads(task_id: str, files: list[UploadFile]) -> list[dict[str, Any]]:
    if not files:
        raise FileIngestionError("Nenhum arquivo enviado.")
    if len(files) > 8:
        raise FileIngestionError("Envie no maximo 8 arquivos por vez.")
    directory = uploads_dir(task_id)
    directory.mkdir(parents=True, exist_ok=True)
    analyses: list[dict[str, Any]] = []
    for upload in files:
        filename = sanitize_upload_filename(upload.filename)
        data = await upload.read()
        if not data:
            raise FileIngestionError(f"Arquivo vazio: {filename}")
        max_size = MAX_ZIP_UPLOAD_BYTES if Path(filename).suffix.lower() == ".zip" else MAX_DOCUMENT_UPLOAD_BYTES
        if len(data) > max_size:
            if Path(filename).suffix.lower() == ".zip":
                raise FileIngestionError(f"ZIP maior que 100 MB: {filename}")
            raise FileIngestionError(f"Arquivo maior que 25 MB: {filename}")
        target = _unique_path(directory, filename)
        target.write_bytes(data)
        try:
            if target.suffix.lower() == ".zip":
                analyses.extend(_analyze_zip(task_id, target))
            else:
                analyses.append(analyze_uploaded_document(task_id, target))
        except FileIngestionError:
            raise
        except Exception as exc:
            raise FileIngestionError(f"Nao foi possivel analisar {filename}: {type(exc).__name__}") from exc
    return analyses


def public_uploaded_file_payloads(analyses: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "path": item["path"],
            "name": item["name"],
            "extension": item["extension"],
            "kind": item.get("kind") or "document",
            "size_bytes": item.get("size_bytes", 0),
            "summary": item.get("summary", ""),
            "archive_contents": item.get("archive_contents", []),
            "geospatial_layers": item.get("geospatial_layers", []),
        }
        for item in analyses
    ]


def build_file_agent_prompt(question: str, analyses: list[dict[str, Any]]) -> str:
    has_geospatial_layer = any(item.get("kind") == "shapefile" or item.get("geospatial_layers") for item in analyses)
    visible_question = question.strip() or "Analise os arquivos enviados."
    if has_geospatial_layer and (geospatial_delivery_from_text(visible_question) or GEOSPATIAL_EDIT_HINT_RE.search(visible_question)):
        visible_question += (
            "\n\nENTREGA_OBRIGATORIA: gere um arquivo .zip final em outputs/ contendo todos os componentes do shapefile "
            "editado/convertido (.shp, .shx, .dbf e .prj/.cpg quando aplicavel)."
        )
    parts = [
        visible_question,
        "",
        "ARQUIVOS_ENVIADOS_PELO_USUARIO_VORTAX:",
    ]
    for item in analyses:
        parts.append(
            f"- {item['path']} ({item.get('extension')}, {item.get('size_bytes')} bytes): {item.get('summary') or 'sem resumo'}"
        )
        extracted = _clip(str(item.get("extracted_text") or ""), PROMPT_PREVIEW_CHARS)
        if extracted:
            parts.append("  CONTEUDO_EXTRAIDO_AMOSTRA:\n" + "\n".join(f"  {line}" for line in extracted.splitlines()))
        if item.get("archive_contents"):
            contents = item.get("archive_contents") or []
            parts.append("  CONTEUDO_DO_ZIP:")
            for content in contents[:80]:
                parts.append(
                    f"  - {content.get('path')} ({content.get('extension') or 'sem extensao'}, {content.get('size_bytes', 0)} bytes)"
                )
        if item.get("geospatial_layers"):
            parts.append("  CAMADAS_GEOESPACIAIS_DETECTADAS:")
            for layer in (item.get("geospatial_layers") or [])[:20]:
                details = layer.get("details") if isinstance(layer.get("details"), dict) else {}
                parts.append(
                    f"  - {layer.get('path')}: {layer.get('summary')} "
                    f"CRS={details.get('crs') or 'ausente'} bounds={details.get('bounds') or 'n/a'}"
                )
    parts.extend(
        [
            "",
            "INSTRUCOES_DE_ARQUIVO_VORTAX:",
            "- Use os arquivos em uploads/ como entrada e fonte de verdade para responder, analisar, editar ou converter.",
            "- Nao sobrescreva arquivos em uploads/. Preserve os originais.",
            "- Se o usuario pedir edicao, correcao, formatacao, conversao ou criacao derivada, gere uma nova versao real em outputs/ com nome descritivo.",
            "- Para XLSX, use openpyxl e valide abrindo o arquivo final.",
            "- Para DOCX, use python-docx e valide abrindo o arquivo final.",
            "- Para CSV, gere UTF-8 com cabecalhos e linhas consistentes.",
            "- Para ZIP, use tambem os arquivos extraidos em uploads/archives/ como entrada; nao trabalhe apenas pelo nome do ZIP.",
            "- Para shapefiles, trate .shp/.shx/.dbf/.prj/.cpg como um conjunto unico; nao entregue apenas .shp isolado.",
            "- Para tarefas geoespaciais, use geopandas, pyogrio, shapely e pyproj. Confira CRS, bounds, atributos, geometrias invalidas e preserve componentes obrigatorios.",
            "- Para alterar tabelas de atributos, preserve geometrias, CRS, contagem/ordem de feicoes e encoding, salvo pedido explicito em contrario.",
            "- Para conversoes GIS, gere outputs em formatos adequados como shapefile completo, GeoJSON, CSV de atributos, XLSX ou relatorio Markdown conforme o pedido.",
            "- Depois de gerar arquivos finais, deixe-os no workspace para o Vortax anexar no chat.",
        ]
    )
    return "\n".join(parts).strip()
